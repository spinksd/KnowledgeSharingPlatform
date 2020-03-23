from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.models import User
from django.core.files.storage import default_storage
from django.core.paginator import Paginator
from django.db.models import Q, Count
from django.http import HttpResponse, FileResponse
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView, FormView
from dal import autocomplete
import docx
from nltk.corpus import stopwords
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx
import os
import RAKE
import re
from taggit.utils import edit_string_for_tags
from .models import Page
from .forms import UploadFileForm, CreateUpdatePageForm, AdvancedSearchForm

# Had to update this to be FormView as opposed to TemplateView such that it would pick up on the form_class attribute
# The AdvancedSearchForm holds the autocomplete widget for the autocomplete contact field in advanced search
class HomePageView(LoginRequiredMixin, FormView):
    template_name = 'website/home.html'
    form_class = AdvancedSearchForm

# Inherits list view as page effectively has a 'list' of latest pages
class PageListView(LoginRequiredMixin, ListView):
    model = Page
    # Update which template this list view interacts with
    template_name = 'website/most_recent.html'
    context_object_name = 'pages'
    # Order pages so latest page is shown at top (This is done by the '-')
    ordering = ['-date_posted']
    # Defining the number of published pages that should be displayed on each page of the list view
    paginate_by = 5

# Inherits list view
class UserPageListView(LoginRequiredMixin, ListView):
    model = Page
    # Update which template this list view interacts with
    template_name = 'website/user_pages.html'
    context_object_name = 'pages'
    # Defining the number of published pages that should be displayed on each page of the list view
    paginate_by = 5

    def get_queryset(self):
        # Get user from User model where username is equal to the username in the url
        # This will 404 if user doesn't exist
        user = get_object_or_404(User, username=self.kwargs.get('username'))
        # Return pages where the user is the author, ordered by date_posted
        return Page.objects.filter(author=user).order_by('-date_posted')

# Inherit from detail view as it provides the 'details' of the published page
class PageDetailView(LoginRequiredMixin, DetailView):
    model = Page
    is_liked = False

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        page = context['page']
        if page.likes.filter(id=self.request.user.id).exists():
            context['is_liked'] = True
        return context

# Inherit from LoginRequiredMixin such that the user is redirected to the login page when attempting to access this page when not logged in
# https://docs.djangoproject.com/en/3.0/topics/auth/default/#the-loginrequired-mixin
# Also inherit from CreateView as this class is for creating a page
class PageCreateView(LoginRequiredMixin, CreateView):
    model = Page
    form_class = CreateUpdatePageForm

    # Get currently logged in user's profile image and description for sidebar
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Check if user has just uploaded a document and set variables appropriately if they have
        # Deletes session variables so they doesn't persist
        # Then if user goes to another page and back to the create page, the previously generated text isn't populated in the create page
        if self.request.session.get('title', None) is not None:
            context['title'] = self.request.session['title']
            del self.request.session['title']
        if self.request.session.get('summary_text', None) is not None:
            context['summary_text'] = self.request.session['summary_text']
            del self.request.session['summary_text']
        if self.request.session.get('tags', None) is not None:
            context['tags'] = self.request.session['tags']
            del self.request.session['tags']
        return context

    # This specifies that the author of the page that is being created is the user who is currently logged in and submits the page creation
    def form_valid(self, form):
        # Specify author as current user (This is the neatest way of setting the author - if author is not specified then Django throws an integrity error as the NOT NULL constraint for author is failed)
        form.instance.author = self.request.user
        # Specify form document
        if self.request.session.get('document_path', None) is not None:
            form.instance.document = self.request.session['document_path']
            del self.request.session['document_path']
        # Call original function to handle the rest of the processing
        return super().form_valid(form)
        

# Inherit from UpdateView as this class is for updating a page
# Also inherit from LoginRequiredMixin with addition of UserPassesTestMixin which allows me to add function to check user is author of the page
# The UserPassTestMixin allows me to add the test_func seen below - where I've added the author check functionality
class PageUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Page
    form_class = CreateUpdatePageForm

    # This specifies that the author of the page that is being created is the user who is currently logged in and submits the page creation
    def form_valid(self, form):
        # Specify author as current user (This is the neatest way of setting the author - if author is not specified then Django throws an integrity error as the NOT NULL constraint for author is failed)
        form.instance.author = self.request.user
        # Call original function to handle the rest of the processing
        return super().form_valid(form)

    def test_func(self):
        # Get current page user is trying to access the edit page for
        page = self.get_object()
        # Compare user who made request against author of the page
        if self.request.user == page.author:
            # If user who made request is the author, return true to allow user to edit the page
            return True
        else:
            return False

# Inherit from delete view as it provides delete functionality of a published page
class PageDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Page
    # Set success_url so that this is where the user is directed to (localhost:8000/) when a page is successfully deleted
    success_url = '/'
    
    # Same test function as above - check user attempting to delete page is the author of the page
    def test_func(self):
        # Get current page user is trying to access the edit page for
        page = self.get_object()
        # Compare user who made request against author of the page
        if self.request.user == page.author:
            # If user who made request is the author, return true to allow user to edit the page
            return True
        else:
            return False

class SearchResultsView(LoginRequiredMixin, ListView):
    model = Page
    context_object_name = 'pages'
    template_name = 'website/search_results.html'
    ordering = ['-likes']
    paginate_by = 5

    # Use GET request for this so that searches can be shared between users
    # Filter on page titles
    def get_queryset(self):
        query = self.request.GET.get('query', None)
        tags = self.request.GET.get('tags', None)
        contacts = self.request.GET.getlist('contacts', None)
        hasDocument = self.request.GET.getlist('hasDocument', None)

        if query is not None:
            # Add first filters checking if the title OR description OR main text of the page contains the text the user's searching for
            filters = Q(title__icontains=query) | Q(summary__icontains=query) | Q(main_text__icontains=query)
        else:
            filters = None

        # If user has speicified any tags, filter pages on those tags
        if tags:
            # Split tags using comma to put them into a list such that the filter filters on all tags individually
            tags = tags.split(',')
            # The below format of '&=' specifies to add an 'AND' Q object to the existing filters
            # It's also possible to use '|=' to specify another 'OR' Q object to further filter on
            filters &= Q(tags__name__in=tags)

        # If user has specified any contacts, filter pages on those contacts
        if contacts:
            # The GET request passes the id's (primary key) of the contacts in as a list of strings
            # The below converts it to a list of integers
            contacts = [int(i) for i in contacts]
            # We then use the contact list of integers (primary keys) to check if any of the pages have that user as a contact
            filters &= Q(contacts__pk__in=contacts)

        # If user has specified page must have a document attachment, filter pages on those that have documents
        if hasDocument:
            # Toggle is returned as an array value.
            # Value is always in position 0 of array, so retrieve text value from array
            hasDocument = hasDocument[0]
            # Check value of whether user wants document in pages for search results and filter accordingly
            if hasDocument == 'no':
                # Document is stored in DB as a char field representing the path to the file, so check if path is not populated
                filters &= Q(document='')
            elif hasDocument == 'yes':
                # The tilde symbol represents a 'negated' query (i.e. a 'NOT'), so this searches for documents that have the path populated
                filters &= ~Q(document='')
        
        # Get pages that match all filters
        pages = Page.objects.filter(filters).distinct()
        # Return matched pages to user
        return pages

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Here I add the search url of the user's current search to the page context
        # I then use this for pagination (see in search_results.html usage of the variable {{ search_url }} )
        # I use regex here to remove the current page number from the string (e.g. '&page=1')
        # Which prevents the url from becoming '&page=1&page=2&page=3' etc.
        context['search_url'] = re.sub('\&page=\d+','', self.request.get_full_path())
        return context

# Inherit from LoginRequiredMixin such that the user is redirected to the login page when attempting to access this page when not logged in
# Also inherit from CreateView as this class is for creating a page
class DocumentUploadView(LoginRequiredMixin, CreateView):
    model = Page
    template_name = 'website/upload_document.html'

    def get(self, request):
        form = UploadFileForm()
        return render(request, 'website/upload_document.html', {'form': form})

    def post(self, request):
        form = UploadFileForm(request.POST, request.FILES)
        # Check form is valid and get the required data from form if it is
        if form.is_valid():
            # Loop through files (assuming only one is uploaded, else this will only parse the last one)
            for filename, file in request.FILES.items():
                document = request.FILES[filename]
            
            # Define location to save the folder
            save_folder = default_storage.location + '/documents/'
            # Check save folder exists, else create if not (if folder doesn't exist, receive an error of suspicious document activity)
            if not os.path.exists(save_folder):
                os.makedirs(save_folder)
            # Set full path of document
            save_path = save_folder + document.name
            # Save the document to the filesystem
            default_storage.save(save_path, document)

            # Generate data from uploaded document
            title = re.sub('\.\w+', '', os.path.basename(save_path))
            self.generate_document_data(request, save_path)
            if request.session['error'] is True:
                default_storage.delete(save_path)
                messages.add_message(self.request, messages.WARNING, 'Invalid file type! Only word documents (.docx) and text files (.txt) can be uploaded.')
                return redirect('upload-document')
            # Add title and document path to session for create-page (can't save document as session variable as it's in memory at this point and is unserialisable)
            request.session['title'] = title
            request.session['document_path'] = save_path
            # Direct user to create-page having set session variables with generated data
            return redirect('create-page')

    def parseDocument(self, filePath):
        fullText = []
        filename, file_extension = os.path.splitext(filePath)
        if file_extension == '.docx':
            document = docx.Document(filePath)
            for paragraph in document.paragraphs:
                fullText.append(paragraph.text)
                #print('fulltext = ' + str(fullText))
            article = '\n'.join(fullText).split('. ')
            #print('article = ' + str(article))
        #elif file_extension == '.pdf':
            #document = 
        elif file_extension == '.txt':
            with open(filePath, 'r') as file:
                fullText.append(file.read())
            #print('fulltext = ' + str(fullText))
            article = fullText[0].split('. ')
            #print('article = ' + str(article))
        # If file extension is not docx or txt then set article to invalid to pass error processing to calling function
        else:
            article = 'invalid'
        
        #article = '\n'.join(fullText).split('. ')
        return article

    def sentence_similarity(self, sentence1, sentence2, stopwords=None):
        if stopwords is None:
            stopwords = []
        sentence1 = [w.lower() for w in sentence1]
        sentence2 = [w.lower() for w in sentence2]
        all_words = list(set(sentence1 + sentence2))
        vector1 = [0] * len(all_words)
        vector2 = [0] * len(all_words)

        # build the vector for the first sentence
        for w in sentence1:
            if w in stopwords:
                continue
            vector1[all_words.index(w)] += 1

        # build the vector for the second sentence
        for w in sentence2:
            if w in stopwords:
                continue
            vector2[all_words.index(w)] += 1

        return 1 - cosine_distance(vector1, vector2)

    def build_similarity_matrix(self, sentences, stop_words):
        # Create an empty similarity matrix
        similarity_matrix = np.zeros((len(sentences), len(sentences)))
        for idx1 in range(len(sentences)):
            for idx2 in range(len(sentences)):
                if idx1 == idx2: #ignore if both are same sentences
                    continue 
                similarity_matrix[idx1][idx2] = self.sentence_similarity(sentences[idx1], sentences[idx2], stop_words)
        return similarity_matrix

    def generate_text_summary(self, text, top_n=5):
        stop_words = stopwords.words('english')
        summary_text = []
        
        # Step 1 - Split text
        sentences = []
        for sentence in text:
            sentences.append(sentence.replace("[^a-zA-Z]", " ").split(" "))        
   
        # Step 2 - Generate similarity matrix across sentences
        sentence_similarity_matrix = self.build_similarity_matrix(sentences, stop_words)
        
        # Step 3 - Rank sentences in similarity matrix
        sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_matrix)
        scores = nx.pagerank(sentence_similarity_graph)

        # Step 4 - Sort sentences by rank and pick highest ranked sentences
        ranked_sentence = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)       
        for i in range(top_n):
            summary_text.append(" ".join(ranked_sentence[i][1]))

        # Step 5 - Return the summary text
        return ". ".join(summary_text) + "."
        
    def generate_tags(self, text):
        Rake = RAKE.Rake(RAKE.SmartStopList())
        keywords = Rake.run(str(text) , minCharacters = 3, maxWords = 3, minFrequency = 2)
        tags = []
        # Limit to 6 tags
        limit = 6
        for index, keyword in zip(range(limit), keywords):
            tags.append(keyword[0] + ',')
        return tags

    def generate_document_data(self, request, filePath):
        article = self.parseDocument(filePath)
        # Check if article was parsed successfully (whether file extension was valid)
        if article == 'invalid':
            request.session['error'] = True
            return
        # Else document has been successfully parsed - set this so check for session variable doesn't error because variable doesn't exist
        request.session['error'] = False
        # Generate tags
        request.session['tags'] = self.generate_tags(article)

        # Generate document summary
        summary_description_num_sentences = 10
        if len(article) < summary_description_num_sentences:
            summary_description_num_sentences = len(article.split)
        request.session['summary_text'] = self.generate_text_summary(article, summary_description_num_sentences)


# Inherits list view as page effectively has a 'list' of latest pages
class TopRatedListView(LoginRequiredMixin, ListView):
    model = Page
    # Update which template this list view interacts with
    template_name = 'website/top_rated.html'
    context_object_name = 'pages'
    # Defining the number of published pages that should be displayed on each page of the list view
    paginate_by = 5

    # Overwrite default queryset to get the total count of likes for each page, and then order by the most liked pages
    def get_queryset(self):
        return Page.objects.annotate(like_count=Count('likes')).order_by('-like_count')

# Use autocomplete package for contact search in create / edit page
class ContactsAutocomplete(LoginRequiredMixin, autocomplete.Select2QuerySetView):
    def get_queryset(self):
        # Extra layer of security ensuring user is logged in
        if not self.request.user.is_authenticated:
             return User.objects.none()
        
        # Initially get all users that have an account
        qs = User.objects.all()

        # Filter user list with the letter(s) entered in contact search on create/edit page
        if self.q:
            qs = qs.filter(username__istartswith=self.q)
        
        # Return filtered set of users
        return qs

def download_document(request, pk):
    # Get page object
    page = get_object_or_404(Page, id=pk)
    response = FileResponse(open(page.document.name, 'rb'), as_attachment=True)
    return response

def like_page(request, pk):
    # Get page object
    page = get_object_or_404(Page, id=pk)
    # Check if user has already liked the page. If so and they've clicked the button again - they're removing their like.
    if page.likes.filter(id=request.user.id).exists():
        page.likes.remove(request.user)
        is_liked = False
    # Else if user has not already liked the page, add a like from the user!
    else:
        page.likes.add(request.user)
        is_liked = True
    # Redirect user to the same published page
    return redirect('published-page', pk)

def about(request):
    return render(request, 'website/about.html', {'title': 'About'})